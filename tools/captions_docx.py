"""Generate a phone-friendly captions.docx in every post's output folder.

Reads the series JSONs (source of truth), applies caption best practices
(research 2026-07-28):
  - IG: hook line first (visible 125 chars), line breaks between sentences,
    3-5 niche hashtags on their own final line.
  - TikTok: keyword front-loaded, generic tags (#parati/#fyp/#viral) stripped,
    3-5 niche hashtags.
  - Facebook: conversational, question-led, NO hashtags.

Writes:  brands/<b>/out/<series>/<slug>/captions.docx  (+ refreshed captions.txt)
         brands/<b>/GUIA-DE-SUBIDA.docx  (one general upload guide per brand)
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1] / "brands"

GENERIC_TAGS = {"parati", "fyp", "foryou", "foryoupage", "viral", "xyzbca"}

ACCENT = {"comehometag": RGBColor(0xB8, 0x5C, 0x1E),
          "qolca": RGBColor(0x08, 0x91, 0xB2),
          "propaga": RGBColor(0xF2, 0x64, 0x3F)}
GRAY = RGBColor(0x66, 0x66, 0x66)


def split_hashtags(text: str) -> tuple[str, list[str]]:
    tags = re.findall(r"#[\wáéíóúñÁÉÍÓÚÑ]+", text)
    body = re.sub(r"\s*#[\wáéíóúñÁÉÍÓÚÑ]+", "", text).strip()
    return body, tags


def fmt_ig(caption: str) -> str:
    body, tags = split_hashtags(caption)
    m = re.match(r"(.+?[.!?…])\s+(.*)", body, re.S)
    hook, rest = (m.group(1), m.group(2)) if m else (body, "")
    parts = [hook]
    if rest:
        parts.append(rest)
    if tags:
        parts.append(" ".join(tags[:5]))
    return "\n\n".join(parts)


def fmt_tt(caption: str) -> str:
    body, tags = split_hashtags(caption)
    tags = [t for t in tags if t.lstrip("#").lower() not in GENERIC_TAGS][:5]
    return body + ("\n" + " ".join(tags) if tags else "")


def fmt_fb(caption: str) -> str:
    body, _tags = split_hashtags(caption)  # FB: hashtags do nothing — drop them
    return body


def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_heading(doc: Document, text: str, color: RGBColor, size: int = 15) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.space_after = Pt(2)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


def add_copybox(doc: Document, text: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade(cell, "F4F4F0")
    cell.paragraphs[0].text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
    doc.add_paragraph().space_after = Pt(4)


def post_docx(brand: str, series: str, post: dict, out: Path) -> None:
    cap = post.get("caption", {})
    acc = ACCENT[brand]
    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run(post["slides"][0].get("h", post["slug"]))
    r.bold = True
    r.font.size = Pt(17)
    sub = doc.add_paragraph()
    r = sub.add_run(f"{brand} · {series} · {post['slug']} · {len(post['slides'])} láminas")
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    if cap.get("tt"):
        add_heading(doc, "TIKTOK", acc)
        add_note(doc, "Fotos de la carpeta tt\\ (9:16) · ⚠ AGREGA UN SONIDO EN TENDENCIA (obligatorio: sin sonido no hay alcance) · pega este caption:")
        add_copybox(doc, fmt_tt(cap["tt"]))

    if cap.get("ig"):
        add_heading(doc, "INSTAGRAM", acc)
        add_note(doc, "Fotos de la carpeta ig\\ (4:5), en orden 01→0" + str(min(len(post['slides']), 9)) + " · pega este caption · si puedes, agrega el ALT de abajo a la lámina 1:")
        add_copybox(doc, fmt_ig(cap["ig"]))
        if post.get("alt"):
            add_note(doc, "ALT (accesibilidad + SEO — en Opciones avanzadas → Escribir texto alternativo):")
            add_copybox(doc, post["alt"])

    if cap.get("fb"):
        add_heading(doc, "FACEBOOK", acc)
        add_note(doc, "Mismas fotos de ig\\ · SIN hashtags (en FB no ayudan) · pega este caption:")
        add_copybox(doc, fmt_fb(cap["fb"]))

    if (out / "li.pdf").exists():
        add_heading(doc, "LINKEDIN", acc)
        add_note(doc, "Sube el archivo li.pdf como DOCUMENTO (no como imágenes) · usa el caption de Instagram sin hashtags, o el de Facebook.")

    doc.save(out / "captions.docx")

    # refresh captions.txt with the same cleaned versions
    txt = []
    if cap.get("tt"):
        txt.append("== TIKTOK (agrega sonido en tendencia) ==\n" + fmt_tt(cap["tt"]))
    if cap.get("ig"):
        txt.append("== INSTAGRAM ==\n" + fmt_ig(cap["ig"]))
        if post.get("alt"):
            txt.append("== ALT (Instagram) ==\n" + post["alt"])
    if cap.get("fb"):
        txt.append("== FACEBOOK (sin hashtags) ==\n" + fmt_fb(cap["fb"]))
    (out / "captions.txt").write_text("\n\n".join(txt), encoding="utf-8")


def guide_docx(brand_dir: Path) -> None:
    acc = ACCENT[brand_dir.name]
    doc = Document()
    r = doc.add_paragraph().add_run("GUÍA DE SUBIDA — " + brand_dir.name.upper())
    r.bold = True
    r.font.size = Pt(18)

    sections = [
        ("CÓMO ESTÁ ORGANIZADO", [
            "Cada carpeta out\\<serie>\\<post> tiene: ig\\ (láminas 4:5 para Instagram y Facebook), tt\\ (láminas 9:16 para TikTok), captions.docx (este formato, con los textos listos para copiar) y captions.txt (lo mismo en texto plano).",
            "Sube las láminas SIEMPRE en orden: 01, 02, 03... La lámina 01 es la portada.",
        ]),
        ("INSTAGRAM", [
            "Formato: las 8 láminas de ig\\ como carrusel.",
            "El caption ya viene con el gancho en la primera línea (los primeros 125 caracteres son los que se ven antes del \"más\") y 3-5 hashtags al final. Pégalo tal cual.",
            "Si puedes: agrega el texto ALT a la lámina 1 (Opciones avanzadas → Texto alternativo). Ayuda al SEO.",
            "Mejor momento: cuando tengas 30 minutos libres DESPUÉS para responder comentarios y DMs — la primera hora decide la distribución.",
        ]),
        ("TIKTOK", [
            "Formato: las láminas de tt\\ como publicación de fotos (carrusel).",
            "⚠ LO MÁS IMPORTANTE: agrega un SONIDO EN TENDENCIA antes de publicar. En TikTok el audio es parte de la distribución incluso en fotos — sin sonido, el post nace invisible.",
            "El caption ya viene corto, con la palabra clave adelante y sin tags genéricos (#parati no ayuda desde 2025).",
        ]),
        ("FACEBOOK", [
            "Formato: las mismas láminas de ig\\.",
            "El caption de FB es distinto a propósito: conversacional y con pregunta (los comentarios son la señal #1 en FB). SIN hashtags — en Facebook no hacen nada.",
            "Responde los primeros comentarios: la conversación llama conversación.",
        ]),
        ("REGLAS GENERALES", [
            "1 carrusel por marca por día. Constancia > volumen.",
            "No repitas el mismo caption en dos redes: ya vienen adaptados, usa cada uno en la suya.",
            "Métricas que importan (revisar el primer lunes de cada mes): ENVÍOS y GUARDADOS en IG, comentarios en FB. Los likes no cuentan la historia.",
            "El post que funcione mejor: avísale a Claude para clonar su formato en la siguiente tanda.",
        ]),
    ]
    if brand_dir.name == "qolca":
        sections.insert(4, ("LINKEDIN (solo Qolca)", [
            "Sube li.pdf como DOCUMENTO en el perfil personal de Carlos (no la página de empresa).",
            "Caption: el de Instagram sin hashtags, o el de Facebook. LinkedIn premia la primera línea con gancho.",
        ]))
    for head, bullets in sections:
        add_heading(doc, head, acc, 13)
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(b)
            r.font.size = Pt(10.5)
    doc.save(brand_dir / "GUIA-DE-SUBIDA.docx")


def main() -> None:
    total = 0
    for bdir in sorted(ROOT.iterdir()):
        posts_dir = bdir / "posts"
        if not posts_dir.is_dir():
            continue
        guide_docx(bdir)
        for sf in sorted(posts_dir.glob("*.json")):
            data = json.loads(sf.read_text(encoding="utf-8"))
            for post in data["posts"]:
                out = bdir / "out" / sf.stem / post["slug"]
                if out.is_dir():
                    post_docx(bdir.name, sf.stem, post, out)
                    total += 1
        print(f"{bdir.name}: guía + docx por post OK")
    print(f"TOTAL: {total} captions.docx")


if __name__ == "__main__":
    main()
